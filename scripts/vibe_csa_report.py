#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
 自动化代码审计报告生成器
基于 vibe-csa JSON 输出生成专业 Word 安全审计报告

Usage:
    python vibe_csa_report.py -i vibe-csa.json -o report.docx [-l logo.png]
"""

import json
import os
import sys
import argparse
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib.pyplot as plt

COLORS = {
    'primary_dark': '#1a3a5c',
    'primary_cyan': '#4ecdc4',
    'accent_teal': '#2a9d8f',
    'text_dark': '#2c3e50',
    'text_gray': '#5a6c7d',
    'bg_light': '#f8f9fa',
    'critical': '#e74c3c',
    'high': '#e67e22',
    'medium': '#f39c12',
    'low': '#3498db',
    'info': '#95a5a6',
}

SEVERITY_COLORS = {
    'critical': '#e74c3c',
    'high': '#e67e22',
    'medium': '#f39c12',
    'low': '#3498db',
}

SEVERITY_LABELS = {
    'critical': '严重',
    'high': '高危',
    'medium': '中危',
    'low': '低危',
}

# =============================================================================
# 攻击路径评分字段映射表
# =============================================================================
AUTH_REQUIRED_MAP = {
    0: '无需登录',
    1: '普通用户',
    2: '特定角色',
    3: '管理员',
}

REQUEST_COMPLEXITY_MAP = {
    0: '单请求',
    1: '需要特定参数',
    2: '多步骤',
    3: '需要特殊时序',
}

SOCIAL_ENGINEERING_MAP = {
    0: '无需',
    1: '需要诱导点击',
    2: '需要特定用户操作',
    3: '高度依赖',
}

EXPLOIT_BARRIER_MAP = {
    0: '无屏障',
    1: '需要特定条件',
    2: '有一定防护',
    3: '有强防护',
}

LEVEL_MAP = {
    'P0': '极易 (0-3分)',
    'P1': '较易 (4-6分)',
    'P2': '有条件 (7-9分)',
    'P3': '困难 (10-12分)',
}


# =============================================================================
# 工具函数
# =============================================================================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml('<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color))
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_custom(doc, text, level=1, color=None):
    """添加自定义标题"""
    if color is None:
        color = COLORS['primary_dark']
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(
            int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        )
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.size = Pt(20)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(14)
            run.font.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph_custom(doc, text, bold=False, color=None, size=Pt(10.5), align='left'):
    """添加自定义段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(
            int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        )
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    return p


def add_code_block(doc, code_text, language='java'):
    """添加代码块样式（浅灰背景 + Consolas 字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = parse_xml('<w:shd {} w:fill="F5F5F5" w:val="clear"/>'.format(nsdecls("w")))
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(44, 62, 80)
    return p


def format_http_request(request_data):
    """
    格式化 HTTP 请求为可直接复制的完整文本。
    动态检测 body / params / headers 字段，支持 form/json 等多种格式。
    """
    lines = []
    method = request_data.get('method', 'GET')
    url = request_data.get('url', '')
    lines.append(f"{method} {url} HTTP/1.1")

    for header, value in request_data.get('headers', {}).items():
        lines.append(f"{header}: {value}")

    body = request_data.get('body')
    params = request_data.get('params')

    if body is not None:
        lines.append("")
        if isinstance(body, dict):
            # 检测 Content-Type 决定输出格式
            content_type = ''
            for h, v in request_data.get('headers', {}).items():
                if h.lower() == 'content-type':
                    content_type = v.lower()
                    break
            if 'application/x-www-form-urlencoded' in content_type:
                form_items = [f"{k}={v}" for k, v in body.items()]
                lines.append('&'.join(form_items))
            else:
                lines.append(json.dumps(body, ensure_ascii=False, indent=2))
        else:
            lines.append(str(body))
    elif params is not None and len(params) > 0:
        lines.append("")
        param_str = '&'.join([f"{k}={v}" for k, v in params.items()])
        lines.append(f"?{param_str}")

    return '\n'.join(lines)


def format_http_response(response_data):
    """格式化 HTTP 响应为可直接复制的完整文本"""
    lines = []
    status = response_data.get('status', 200)
    status_text = {
        200: 'OK', 201: 'Created', 204: 'No Content',
        301: 'Moved Permanently', 302: 'Found', 304: 'Not Modified',
        400: 'Bad Request', 401: 'Unauthorized', 403: 'Forbidden',
        404: 'Not Found', 500: 'Internal Server Error', 502: 'Bad Gateway'
    }.get(status, 'OK')
    lines.append(f"HTTP/1.1 {status} {status_text}")

    for header, value in response_data.get('headers', {}).items():
        lines.append(f"{header}: {value}")

    body = response_data.get('body')
    if body is not None:
        lines.append("")
        if isinstance(body, dict):
            lines.append(json.dumps(body, ensure_ascii=False, indent=2))
        else:
            lines.append(str(body))

    return '\n'.join(lines)


# =============================================================================
# 图表生成
# =============================================================================

def generate_charts(data, output_dir):
    """生成统计图表并返回文件路径列表"""
    charts = []
    summary = data['audit']['summary']
    findings = data['findings']

    plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

    # 1. 漏洞严重程度分布饼图
    fig, ax = plt.subplots(figsize=(8, 6))
    severity_counts = {
        '严重': summary['critical'],
        '高危': summary['high'],
        '中危': summary['medium'],
        '低危': summary['low']
    }
    severity_counts = {k: v for k, v in severity_counts.items() if v > 0}
    if severity_counts:
        colors_pie = []
        for k in severity_counts.keys():
            colors_pie.append({
                '严重': COLORS['critical'], '高危': COLORS['high'],
                '中危': COLORS['medium'], '低危': COLORS['low']
            }[k])
        wedges, texts, autotexts = ax.pie(
            severity_counts.values(),
            labels=severity_counts.keys(),
            autopct='%1.1f%%',
            colors=colors_pie,
            startangle=90,
            textprops={'fontsize': 12, 'fontweight': 'bold'}
        )
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontsize(11)
    ax.set_title('漏洞严重程度分布', fontsize=16, fontweight='bold',
                 color=COLORS['primary_dark'], pad=20)
    plt.tight_layout()
    chart1_path = os.path.join(output_dir, 'chart_severity_pie.png')
    plt.savefig(chart1_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    charts.append(chart1_path)


    # 3. 漏洞类型分布横向柱状图
    fig, ax = plt.subplots(figsize=(10, 5))
    vuln_types = {}
    for f in findings:
        title = f['title']
        vuln_type = title.split('——')[0].strip() if '——' in title else title
        vuln_types[vuln_type] = vuln_types.get(vuln_type, 0) + 1

    y_pos = range(len(vuln_types))
    bars = ax.barh(y_pos, vuln_types.values(), color=COLORS['primary_cyan'])
    ax.set_yticks(y_pos)
    ax.set_yticklabels(vuln_types.keys(), fontsize=11)
    ax.set_xlabel('数量', fontsize=12, fontweight='bold')
    ax.set_title('漏洞类型分布', fontsize=16, fontweight='bold',
                 color=COLORS['primary_dark'], pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    for i, bar in enumerate(bars):
        width = bar.get_width()
        if width > 0:
            ax.text(width, bar.get_y() + bar.get_height()/2.,
                    f' {int(width)}', ha='left', va='center', fontsize=11, fontweight='bold')
    plt.tight_layout()
    chart3_path = os.path.join(output_dir, 'chart_type_bar.png')
    plt.savefig(chart3_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    charts.append(chart3_path)

    return charts


# =============================================================================
# 报告生成主函数
# =============================================================================

def get_finding_status(finding, audit_stage):
    """
    根据动态验证结果确定漏洞状态
    
    Args:
        finding: 漏洞数据字典
        audit_stage: 审计阶段 (static_audit 或 dynamic_verification)
    
    Returns:
        str: 状态标签（已确认/待验证/验证未成功）
    """
    if audit_stage != 'dynamic_verification':
        return finding.get('status', 'CONFIRMED')
    
    # 动态验证阶段，根据 poc.result 确定状态（忽略大小写）
    poc = finding.get('poc', {})
    poc_result = poc.get('result', 'pending').lower()
    
    # 状态映射逻辑：根据 poc.result 值判断（忽略大小写）
    if poc_result in ('success', 'confirmed'):
        return 'CONFIRMED'  # 已确认
    elif poc_result == 'pending':
        return 'HYPOTHESIS'  # 待验证
    else:
        # 其他状态（failure、skipped 等）归类为验证未成功
        return 'FAILED'  # 验证未成功


def generate_report(data, logo_path, output_path):
    """
    生成 Word 安全审计报告

    Args:
        data: 审计 JSON 数据 (dict)
        logo_path: Logo 图片路径 (str)
        output_path: 输出 Word 文件路径 (str)
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(44, 62, 80)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 创建临时目录存放图表
    temp_dir = os.path.join(os.path.dirname(output_path) or '.', '_temp_charts')
    os.makedirs(temp_dir, exist_ok=True)

    # =====================================================================
    # 封面
    # =====================================================================
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2.5))

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data['audit']['title'])
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(
        int(COLORS['primary_dark'][1:3], 16),
        int(COLORS['primary_dark'][3:5], 16),
        int(COLORS['primary_dark'][5:7], 16)
    )
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    title.paragraph_format.space_after = Pt(20)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('自动化代码安全审计报告')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(
        int(COLORS['primary_cyan'][1:3], 16),
        int(COLORS['primary_cyan'][3:5], 16),
        int(COLORS['primary_cyan'][5:7], 16)
    )
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    subtitle.paragraph_format.space_after = Pt(40)

    # 基本信息表格
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'

    info_data = [
        ('审计编号', data['audit']['audit_id']),
        ('审计对象', data['audit']['repository']),
        ('审计周期', f"{data['audit']['audit_date']['start']} 至 {data['audit']['audit_date']['end']}"),
        ('目标语言', ', '.join(data['audit']['language'])),
        ('报告生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
    ]

    for i, (key, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row.cells[0], COLORS['primary_dark'][1:])
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =====================================================================
    # 执行摘要
    # =====================================================================
    add_heading_custom(doc, '一、执行摘要', level=1)

    summary = data['audit']['summary']
    findings = data['findings']
    audit_stage = data['audit'].get('stage', 'static_audit')
    
    # 动态统计三种状态的数量
    confirmed_count = 0  # 已确认
    hypothesis_count = 0  # 待验证
    failed_count = 0  # 验证未成功
    
    for finding in findings:
        status = get_finding_status(finding, audit_stage)
        if status == 'CONFIRMED':
            confirmed_count += 1
        elif status == 'HYPOTHESIS':
            hypothesis_count += 1
        elif status == 'FAILED':
            failed_count += 1
    
    add_paragraph_custom(doc,
        f"本次安全审计共发现 {summary['total']} 个安全漏洞，其中严重漏洞 {summary['critical']} 个、"
        f"高危漏洞 {summary['high']} 个、中危漏洞 {summary['medium']} 个、低危漏洞 {summary['low']} 个。"
        f"已确认 {confirmed_count} 个，待验证 {hypothesis_count} 个，验证未成功 {failed_count} 个。")

    # 风险等级统计表格
    add_heading_custom(doc, '1.1 风险等级统计', level=2, color=COLORS['text_gray'])

    stat_table = doc.add_table(rows=5, cols=4)
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stat_table.style = 'Table Grid'

    headers = ['风险等级', '数量', '占比', '状态']
    header_row = stat_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLORS['primary_dark'][1:])

    severity_data = [
        ('严重', summary['critical'],
         f"{summary['critical']/summary['total']*100:.1f}%" if summary['total'] > 0 else '0%', '需立即修复'),
        ('高危', summary['high'],
         f"{summary['high']/summary['total']*100:.1f}%" if summary['total'] > 0 else '0%', '建议优先修复'),
        ('中危', summary['medium'],
         f"{summary['medium']/summary['total']*100:.1f}%" if summary['total'] > 0 else '0%', '计划修复'),
        ('低危', summary['low'],
         f"{summary['low']/summary['total']*100:.1f}%" if summary['total'] > 0 else '0%', '酌情修复'),
    ]
    severity_colors_list = [COLORS['critical'], COLORS['high'], COLORS['medium'], COLORS['low']]

    for i, (level, count, ratio, status) in enumerate(severity_data):
        row = stat_table.rows[i + 1]
        row.cells[0].text = level
        row.cells[1].text = str(count)
        row.cells[2].text = ratio
        row.cells[3].text = status
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row.cells[0], severity_colors_list[i][1:])
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    # =====================================================================
    # 图表展示
    # =====================================================================
    add_heading_custom(doc, '1.2 统计图表', level=2, color=COLORS['text_gray'])

    charts = generate_charts(data, temp_dir)

    for chart_path in charts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(5.5))
        p.paragraph_format.space_after = Pt(12)

    # 漏洞列表（含审计方法）
    add_heading_custom(doc, '1.3 漏洞列表', level=2, color=COLORS['text_gray'])

    findings = data['findings']
    audit_stage = data['audit'].get('stage', 'static_audit')
    audit_method = '动态漏洞验证' if audit_stage == 'dynamic_verification' else '静态代码审计'
    
    findings_table = doc.add_table(rows=1 + len(findings), cols=5)
    findings_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    findings_table.style = 'Table Grid'

    findings_headers = ['序号', '漏洞名称', '严重程度', '状态', '审计方法']
    for i, header in enumerate(findings_headers):
        cell = findings_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLORS['primary_dark'][1:])

    for idx, finding in enumerate(findings, 1):
        row = findings_table.rows[idx]
        
        severity = finding['severity']
        severity_label = SEVERITY_LABELS.get(severity, severity)
        severity_color = SEVERITY_COLORS.get(severity, COLORS['info'])
        
        # 根据审计阶段和实际 result 确定状态
        status = get_finding_status(finding, audit_stage)
        status_labels = {
            'CONFIRMED': '已确认',
            'HYPOTHESIS': '待验证',
            'FAILED': '验证未成功'
        }
        status_label = status_labels.get(status, status)
        status_colors_map = {
            'CONFIRMED': COLORS['critical'],  # 红色
            'HYPOTHESIS': COLORS['info'],     # 灰色
            'FAILED': '#808080'               # 灰色
        }
        status_color = status_colors_map.get(status, COLORS['info'])

        # 根据 poc.result 动态确定审计方法（忽略大小写）
        poc_result = finding.get('poc', {}).get('result', '').lower()
        if poc_result == 'pending':
            finding_audit_method = '静态代码审计'
        else:
            finding_audit_method = audit_method

        row.cells[0].text = str(idx)
        row.cells[1].text = finding['title']
        row.cells[2].text = severity_label
        row.cells[3].text = status_label
        row.cells[4].text = finding_audit_method
        
        for cell_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                if cell_idx == 0 or cell_idx == 3 or cell_idx == 4:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        set_cell_shading(row.cells[2], severity_color[1:])
        for paragraph in row.cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
        
        # 状态列背景色
        set_cell_shading(row.cells[3], status_color[1:])
        for paragraph in row.cells[3].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
        
        # 审计方法列背景色
        if finding_audit_method == '静态代码审计':
            set_cell_shading(row.cells[4], COLORS['accent_teal'][1:])
        else:
            set_cell_shading(row.cells[4], COLORS['primary_cyan'][1:])
        for paragraph in row.cells[4].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_page_break()

    # =====================================================================
    # 漏洞详情
    # =====================================================================
    add_heading_custom(doc, '二、漏洞详情', level=1)

    findings = data['findings']
    audit_stage = data['audit'].get('stage', 'static_audit')
    for idx, finding in enumerate(findings, 1):
        add_heading_custom(doc, f"2.{idx} {finding['title']}", level=2)

        # 严重级别标签 + 基本信息
        info_para = doc.add_paragraph()
        info_para.paragraph_format.space_after = Pt(12)

        severity = finding['severity']
        severity_label = SEVERITY_LABELS.get(severity, severity)
        severity_color = SEVERITY_COLORS.get(severity, COLORS['info'])
        
        # 根据审计阶段和实际 result 确定状态
        status = get_finding_status(finding, audit_stage)
        status_labels = {
            'CONFIRMED': '已确认',
            'HYPOTHESIS': '待验证',
            'FAILED': '验证未成功'
        }
        status_label = status_labels.get(status, status)

        run = info_para.add_run(f"  {severity_label}  ")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        shading = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls("w"), severity_color[1:]))
        run._element.get_or_add_rPr().append(shading)

        run = info_para.add_run(
            f"    漏洞编号: {finding['vuln_id']}    "
            f"DKTSS评分: {finding.get('dktss_score', 'N/A')}    "
            f"状态: {status_label}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(
            int(COLORS['text_gray'][1:3], 16),
            int(COLORS['text_gray'][3:5], 16),
            int(COLORS['text_gray'][5:7], 16)
        )
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # 漏洞描述
        add_heading_custom(doc, '漏洞描述', level=3, color=COLORS['accent_teal'])
        add_paragraph_custom(doc, finding['description'])

        # 漏洞位置
        add_heading_custom(doc, '漏洞位置', level=3, color=COLORS['accent_teal'])
        loc = finding['location']
        add_paragraph_custom(doc, f"文件: {loc['file']}")
        if loc.get('function'):
            add_paragraph_custom(doc, f"函数: {loc['function']} (第 {loc['line_start']}-{loc['line_end']} 行)")
        else:
            add_paragraph_custom(doc, f"位置: 第 {loc['line_start']}-{loc['line_end']} 行")
        if loc.get('route'):
            add_paragraph_custom(doc, f"路由: {loc['route']}")

        add_paragraph_custom(doc, '漏洞代码片段:', bold=True, color=COLORS['text_gray'])
        add_code_block(doc, loc['snippet'])

        # 数据流追踪（新结构在 analysis.data_flow 中）
        data_flow = finding.get('analysis', {}).get('data_flow', [])
        if data_flow:
            add_heading_custom(doc, '数据流追踪', level=3, color=COLORS['accent_teal'])
            for flow in data_flow:
                flow_type_label = {'source': '数据源', 'propagation': '传播点', 'validation': '验证点', 'sink': '执行点'}
                label = flow_type_label.get(flow['type'], flow['type'])
                add_paragraph_custom(doc, f"步骤 {flow['step']} [{label}] {flow['location']}")
                add_paragraph_custom(doc, f"  → {flow['desc']}", color=COLORS['text_gray'])
                if flow.get('code'):
                    add_code_block(doc, flow['code'])

        # 安全控制分析
        security_controls = finding.get('analysis', {}).get('security_controls', [])
        if security_controls:
            add_heading_custom(doc, '安全控制分析', level=3, color=COLORS['accent_teal'])
            for control in security_controls:
                add_paragraph_custom(doc, f"• {control.get('control', 'N/A')}", bold=True)
                add_paragraph_custom(doc, f"  位置: {control.get('location', 'N/A')}", color=COLORS['text_gray'])
                assessment_labels = {
                    'present_effective': '有效',
                    'present_bypassable': '可绕过',
                    'absent': '缺失'
                }
                assessment_label = assessment_labels.get(control.get('assessment', ''), control.get('assessment', 'N/A'))
                add_paragraph_custom(doc, f"  评估: {assessment_label}", color=COLORS['text_gray'])
                if control.get('bypass_notes'):
                    add_paragraph_custom(doc, f"  绕过说明: {control['bypass_notes']}", color=COLORS['high'])

        # 绕过策略
        bypass_strategy = finding.get('analysis', {}).get('bypass_strategy', {})
        if bypass_strategy.get('feasible'):
            add_heading_custom(doc, '绕过分析', level=3, color=COLORS['accent_teal'])
            difficulty_labels = {'low': '低', 'medium': '中', 'high': '高'}
            difficulty = difficulty_labels.get(bypass_strategy.get('difficulty', ''), bypass_strategy.get('difficulty', 'N/A'))
            add_paragraph_custom(doc, f"绕过可行性: {'可行' if bypass_strategy.get('feasible') else '不可行'}")
            add_paragraph_custom(doc, f"绕过难度: {difficulty}")
            
            if bypass_strategy.get('ideas'):
                add_paragraph_custom(doc, '绕过思路:', bold=True, color=COLORS['high'])
                for idea in bypass_strategy['ideas']:
                    add_code_block(doc, f"技术: {idea.get('technique', 'N/A')}\nPayload: {idea.get('payload_hint', 'N/A')}\n原因: {idea.get('reason', 'N/A')}")


        # 攻击路径评估（仅在有数据时显示）
        attack = finding.get('attack_path', {})
        if attack:
            add_heading_custom(doc, '攻击路径评估', level=3, color=COLORS['accent_teal'])
            
            attack_info = []
            if 'total' in attack:
                attack_info.append(
                    f"攻击路径综合得分: {attack['total']} (范围 0-12，越低越容易利用)")
            if 'level' in attack:
                level_desc = LEVEL_MAP.get(attack['level'], attack['level'])
                attack_info.append(f"攻击难度等级: {attack['level']} —— {level_desc}")
            if 'auth_required' in attack:
                auth_desc = AUTH_REQUIRED_MAP.get(attack['auth_required'], str(attack['auth_required']))
                attack_info.append(f"认证要求: {attack['auth_required']} —— {auth_desc}")
            if 'request_complexity' in attack:
                req_desc = REQUEST_COMPLEXITY_MAP.get(attack['request_complexity'], str(attack['request_complexity']))
                attack_info.append(f"请求复杂度: {attack['request_complexity']} —— {req_desc}")
            if 'social_engineering' in attack:
                soc_desc = SOCIAL_ENGINEERING_MAP.get(attack['social_engineering'], str(attack['social_engineering']))
                attack_info.append(f"社交工程依赖: {attack['social_engineering']} —— {soc_desc}")
            if 'exploit_barrier' in attack:
                bar_desc = EXPLOIT_BARRIER_MAP.get(attack['exploit_barrier'], str(attack['exploit_barrier']))
                attack_info.append(f"利用屏障: {attack['exploit_barrier']} —— {bar_desc}")

            for info in attack_info:
                add_paragraph_custom(doc, f"• {info}")

        # POC 验证（新结构：dynamic_verification + poc）
        # 如果整体审计方法是静态代码审计，则不显示动态验证部分
        audit_stage = data['audit'].get('stage', 'static_audit')
        if audit_stage != 'static_audit':
            dynamic_verification = finding.get('dynamic_verification', {})
            poc = finding.get('poc', {})
            poc_steps = poc.get('steps', [])
            
            # 只在有动态验证数据或 POC 步骤时显示
            if dynamic_verification.get('state') or poc_steps:
                add_heading_custom(doc, '动态验证', level=3, color=COLORS['accent_teal'])
                
                # 验证状态
                state = dynamic_verification.get('state', 'unknown')
                state_labels = {
                    'verified': '已验证',
                    'not_started': '未开始',
                    'failed': '验证失败',
                    'partial': '部分验证'
                }
                state_label = state_labels.get(state, state)
                state_color = {
                    'verified': COLORS['accent_teal'],
                    'not_started': COLORS['info'],
                    'failed': COLORS['critical'],
                    'partial': COLORS['medium']
                }.get(state, COLORS['text_gray'])
                
                add_paragraph_custom(doc, f"验证状态: {state_label}", bold=True, color=state_color)
                
                # 如果有运行时备注
                if dynamic_verification.get('runtime_notes'):
                    add_paragraph_custom(doc, dynamic_verification['runtime_notes'], color=COLORS['text_gray'])
                
                # 如果有最终证据
                final_evidence = dynamic_verification.get('final_evidence', {})
                if final_evidence.get('summary'):
                    add_paragraph_custom(doc, '验证证据:', bold=True, color=COLORS['primary_dark'])
                    add_paragraph_custom(doc, final_evidence['summary'])
                    
                    # 显示证据片段
                    if final_evidence.get('snippets'):
                        for snippet in final_evidence['snippets']:
                            add_code_block(doc, f"[{snippet.get('signature_type', 'N/A')}] {snippet.get('snippet', '')}")
            
            # POC 步骤详情（HTTP 请求/响应）
            if poc_steps:
                add_heading_custom(doc, 'POC 步骤详情', level=3, color=COLORS['accent_teal'])
                
                for step_idx, step in enumerate(poc_steps, 1):
                    req = step.get('request', {})
                    resp = step.get('response', {})
                    
                    # 步骤名称
                    step_name = step.get('name', f'步骤 {step_idx}')
                    add_paragraph_custom(doc, f"步骤 {step_idx}: {step_name}", bold=True, color=COLORS['primary_dark'])
                    
                    # HTTP 请求
                    if req:
                        add_paragraph_custom(doc, '【HTTP Request】', bold=True, color=COLORS['primary_dark'])
                        add_code_block(doc, format_http_request(req))
                    
                    # HTTP 响应
                    if resp:
                        add_paragraph_custom(doc, '【HTTP Response】', bold=True, color=COLORS['primary_dark'])
                        add_code_block(doc, format_http_response(resp))
                    
                    # 响应中的证据标记
                    if resp.get('_evidence_match'):
                        add_paragraph_custom(doc, '【证据匹配】', bold=True, color=COLORS['accent_teal'])
                        for evidence in resp['_evidence_match']:
                            add_code_block(doc, f"类型: {evidence.get('type', 'N/A')}\n片段: {evidence.get('snippet', '')}\n强度: {evidence.get('strength', 'N/A')}")
                
                # 验证结果总结
                add_paragraph_custom(doc, '【验证结果总结】', bold=True, color=COLORS['primary_dark'])
                result = poc.get('result', 'N/A')
                
                if result == 'success':
                    result_color = COLORS['accent_teal']
                    result_text = 'SUCCESS (验证成功)'
                elif result == 'failure':
                    result_color = COLORS['critical']
                    result_text = 'FAILURE (验证失败)'
                elif result == 'skipped':
                    result_color = COLORS['medium']
                    result_text = 'SKIPPED (源码确认)'
                elif result == 'pending':
                    result_color = COLORS['info']
                    result_text = 'PENDING (待验证)'
                else:
                    result_color = COLORS['info']
                    result_text = result.upper()
                
                add_paragraph_custom(doc, f"结果: {result_text}", color=result_color, bold=True)
                if poc.get('evidence'):
                    add_paragraph_custom(doc, f"证据: {poc['evidence']}")
                
                # 失败日志
                if poc.get('failure_log'):
                    add_paragraph_custom(doc, '失败详情:', bold=True, color=COLORS['critical'])
                    for failure in poc['failure_log']:
                        if isinstance(failure, str):
                            add_paragraph_custom(doc, f"• {failure}")
                        else:
                            add_paragraph_custom(doc, f"• 步骤 {failure.get('step', 'N/A')}: {failure.get('reason', 'N/A')}")
                            if failure.get('hypothesis'):
                                add_paragraph_custom(doc, f"  分析: {failure['hypothesis']}")
                            if failure.get('next_action'):
                                add_paragraph_custom(doc, f"  建议: {failure['next_action']}")
            elif not dynamic_verification.get('state') or dynamic_verification.get('state') == 'not_started':
                # 没有 POC 步骤且未进行动态验证
                pass  # 不显示 POC 章节

        # 修复建议
        add_heading_custom(doc, '修复建议', level=3, color=COLORS['accent_teal'])

        remediation = finding.get('remediation', {})
        if remediation.get('short_term'):
            add_paragraph_custom(doc, '短期缓解措施:', bold=True, color=COLORS['high'])
            add_paragraph_custom(doc, remediation['short_term'])

        if remediation.get('long_term'):
            add_paragraph_custom(doc, '长期修复方案:', bold=True, color=COLORS['accent_teal'])
            add_paragraph_custom(doc, remediation['long_term'])

        # 代码修复示例
        fix = finding.get('fix', {})
        if fix:
            add_heading_custom(doc, '代码修复示例', level=3, color=COLORS['accent_teal'])
            if fix.get('before'):
                add_paragraph_custom(doc, '当前代码片段:', bold=True, color=COLORS['critical'])
                add_code_block(doc, fix['before'])
            if fix.get('after'):
                add_paragraph_custom(doc, '代码整改参考建议:', bold=True, color=COLORS['accent_teal'])
                add_code_block(doc, fix['after'])

        # 分隔线
        if idx < len(findings):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(20)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(
                int(COLORS['primary_cyan'][1:3], 16),
                int(COLORS['primary_cyan'][3:5], 16),
                int(COLORS['primary_cyan'][5:7], 16)
            )
            run.font.size = Pt(8)

    # =====================================================================
    # 附录
    # =====================================================================
    doc.add_page_break()
    add_heading_custom(doc, '附录', level=1)

    add_heading_custom(doc, 'A. 审计方法说明', level=2, color=COLORS['text_gray'])
    add_paragraph_custom(doc,
        '本次审计采用自动化静态代码分析（SAST）结合人工验证的方式进行。'
        '审计范围覆盖目标仓库的全部源代码，重点关注输入验证、认证授权、数据保护、'
        '安全配置等安全维度。')

    add_heading_custom(doc, 'B. 评分标准 (DKTSS)', level=2, color=COLORS['text_gray'])
    add_paragraph_custom(doc,
        'DKTSS (Dynamic Knowledge-based Threat Scoring System) 评分范围为 0-10 分，'
        '评分越高表示风险越大。计算公式：min(10, Base - Friction + Weapon + Ver)，'
        '比 CVSS 更贴近实战，综合考虑了利用摩擦度和武器化程度。')

    score_table = doc.add_table(rows=5, cols=3)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    score_table.style = 'Table Grid'

    score_headers = ['评分区间', '风险等级', '处理优先级']
    for i, header in enumerate(score_headers):
        cell = score_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, COLORS['primary_dark'][1:])

    score_data = [
        ('8.0 - 10.0', '严重', '立即修复'),
        ('6.0 - 7.9', '高危', '优先修复'),
        ('4.0 - 5.9', '中危', '计划修复'),
        ('0.0 - 3.9', '低危', '酌情修复'),
    ]
    for i, (score_range, level, priority) in enumerate(score_data):
        row = score_table.rows[i + 1]
        row.cells[0].text = score_range
        row.cells[1].text = level
        row.cells[2].text = priority
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存文档
    doc.save(output_path)

    # 清理临时文件
    for chart in charts:
        if os.path.exists(chart):
            os.remove(chart)
    if os.path.exists(temp_dir):
        os.rmdir(temp_dir)

    print(f"[OK] Report generated: {os.path.abspath(output_path)}")
    return output_path


# =============================================================================
# 命令行入口
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='自动化代码审计报告生成器',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python vibe_csa_report.py -i vibe-csa.json -o report.docx
  python vibe_csa_report.py -i audit.json -o report.docx -l logo.png
        """
    )

    parser.add_argument(
        '-i', '--input',
        required=True,
        help='输入的 vibe-csa JSON 审计数据文件路径'
    )
    parser.add_argument(
        '-o', '--output',
        required=True,
        help='输出的 Word 报告文件路径 (.docx)'
    )
    parser.add_argument(
        '-l', '--logo',
        default=None,
        help='Logo 图片路径 (可选，用于封面品牌展示)'
    )

    args = parser.parse_args()

    # 验证输入文件
    if not os.path.exists(args.input):
        print(f"[ERROR] Input file not found: {args.input}")
        sys.exit(1)

    # 验证 JSON 格式
    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"[ERROR] JSON parse failed: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] Failed to read file: {e}")
        sys.exit(1)

    # 验证必要字段
    required_keys = ['audit', 'findings']
    for key in required_keys:
        if key not in data:
            print(f"[ERROR] JSON missing required field: {key}")
            sys.exit(1)

    # 确保输出目录存在
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # 生成报告
    try:
        generate_report(data, args.logo, args.output)
    except Exception as e:
        print(f"[ERROR] Failed to generate report: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
